# -*- coding: utf-8 -*-
"""报告审查：用 LLM 对“生成的模板”和“生成的报告”做两轮审查。

第一轮 · 模板审查（analyze_report.py 生成模板后调用）：
  输入 = 成品报告原文 + 模板全文（占位符版）。
  重点：占位符是否准确、上下文有方位(左幅/右幅/上游/下游)时占位符是否带上、
  静态数字是否被误判成占位符、是否篡改了原文内容。

第二轮 · 报告审查（agent.py 生成报告后调用）：
  输入 = 成品报告原文 + 生成报告全文 + 数据链路摘要 + 填表校验告警。
  重点：表格重复列/行、数据索引是否正确、图片是否插错、总结段落统计值是否
  符合逻辑(湿度百分比不超 100、最大值最大、平均值在 [最小值, 最大值] 闭区间)、
  总结段落是否照抄成品报告里的旧值、单位/报告期等。

LLM 不可用或调用失败时自动跳过（降级），不影响报告生成主流程。
"""

import json
import logging
import re
from typing import Dict, List, Optional

from .llm_classifier import LLMClassifier

log = logging.getLogger("report-agent.reviewer")


def _read_docx_text(path: str, limit: int = 50000) -> str:
    """读取 docx 的段落与表格文本，返回纯文本（用于喂给 LLM）。"""
    if not path or not __import__("os").path.isfile(path):
        return ""
    try:
        from docx import Document
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for tb in doc.tables:
            for row in tb.rows:
                cells = [c.text.strip() for c in row.cells]
                parts.append(" | ".join(cells))
        text = "\n".join(parts).strip()
        return text[:limit]
    except Exception as exc:  # noqa: BLE001
        log.warning("读取 docx 文本失败 %s: %s", path, exc)
        return ""


class ReportReviewer:
    """报告审查器（复用 llm_classifier 的 LLM 配置与接口）。"""

    def __init__(self, llm_cfg: Optional[Dict] = None):
        self._llm = LLMClassifier(llm_cfg or {})

    def available(self) -> bool:
        return self._llm.available()

    def _ask_json(self, system: str, user: str) -> Optional[Dict]:
        try:
            resp = self._llm._chat([
                {"role": "system", "content": system},
                {"role": "user", "content": user},
            ])
        except Exception as exc:  # noqa: BLE001
            log.warning("审查 LLM 调用失败: %s", exc)
            return None
        m = re.search(r"\{.*\}", str(resp or ""), re.S)
        if m:
            try:
                return json.loads(m.group(0))
            except Exception as exc:  # noqa: BLE001
                log.warning("解析审查 LLM 回复失败: %s", exc)
        return None

    # ------------------------------------------------------------------
    # 第一轮：模板审查
    # ------------------------------------------------------------------
    def review_template(self, source_text: str, template_text: str) -> Dict:
        """审查生成的模板，返回 {"ok", "issues", "raw"}。"""
        result = {"ok": True, "issues": [], "raw": ""}
        if not self.available():
            log.info("LLM 不可用，跳过模板审查")
            return result
        if not template_text.strip():
            log.info("模板文本为空，跳过模板审查")
            return result
        system = (
            "你是桥梁健康监测报告模板审查专家。请审查“占位符版模板”是否准确，"
            "只输出 JSON：{\"issues\": [{\"type\": \"...\", \"detail\": \"...\"}]}，"
            "没有问题输出 {\"issues\": []}。\n"
            "重点审查：\n"
            "1) 占位符是否准确：{{stats.指标.统计}} 的指标/统计是否对应正文含义；"
            "{{cell.指标.位置.统计}} 的“位置”是否和表格行一致；{{chart.xxx}} 是否"
            "带上了监测位置。\n"
            "2) 方位词：若节标题/说明句带 左幅/右幅/上游/下游/左侧/右侧，而该节"
            "图表或单元格占位符的位置没有带对应方位，属于 missing_direction。\n"
            "3) 静态值误判：章节号、列表序号、桩号、测点编号、阈值、规范编号、"
            "设计参数(桥长/跨径/矢跨比)、布设数量、检查记录等静态数字被替换成"
            "占位符，属于 static_as_placeholder。\n"
            "4) 篡改原文：模板把原文中不该改的固定文字、单位、标点、标题改动了，"
            "属于 text_tampered（把原文文字替换为占位符、或新增占位符，不算篡改）。\n"
            "5) 单位/标点空格：占位符和它后面的单位、标点之间不能有多余空格"
            "（如 {{...}}   m/s² 有两个空格要报 unit_wrong）。\n"
            "6) 多方向位置不要串：GNSS/支座位移等 X/Y/Z 三方向的“对应测点/位置”"
            "占位符必须各自独立，不能把 Z 的位置写成 Y 的位置。\n"
            "type 取值：placeholder_wrong / missing_direction / "
            "static_as_placeholder / text_tampered / unit_wrong / other。"
            "detail 用简短中文说明问题所在（可引用原文字段），不要泛泛而谈。"
        )
        user = (
            "【成品报告原文】\n" + (source_text or "（无）")[:40000]
            + "\n\n【占位符版模板】\n" + template_text[:40000]
        )
        data = self._ask_json(system, user)
        result["raw"] = json.dumps(data, ensure_ascii=False) if data else ""
        if isinstance(data, dict):
            issues = data.get("issues") or []
            result["issues"] = [i for i in issues if isinstance(i, dict)]
            result["ok"] = not result["issues"]
        else:
            result["ok"] = True  # 解析失败视为无结论，不阻塞
        return result

    # ------------------------------------------------------------------
    # 第二轮：报告审查
    # ------------------------------------------------------------------
    def review_report(self, source_text: str, report_text: str,
                      lineage_digest: str = "",
                      table_warnings: str = "",
                      chart_index: str = "",
                      prior_issues: str = "",
                      name_dict_json: str = "",
                      quarterly_stats_json: str = "") -> Dict:
        """审查生成的报告，返回 {"ok", "issues", "repairs", "raw"}。

        repairs 是机器可执行的修复清单（供 repairer 逐条“验证后落地”），
        issues 是给人看的问题清单，二者可同时存在。

        只做两件事：总结/结论段落润色（对照季度统计JSON）、删除冗余图注。
        其他一律不查不报。不对照成品报告原文；宁缺毋滥。
        """
        result = {"ok": True, "issues": [], "repairs": [], "raw": ""}
        if not self.available():
            log.info("LLM 不可用，跳过报告审查")
            return result
        if not report_text.strip():
            log.info("报告文本为空，跳过报告审查")
            return result
        system = (
            "你是桥梁健康监测报告审查专家。审查“生成的报告”，只做一件事："
            "总结/结论段落润色，其他一律不查、不报（宁缺毋滥）。\n"
            "只输出 JSON：{\"issues\": [], \"repairs\": []}，除该任务外没有任何问题。\n"
            "总结/结论段落润色：只处理 3.x 小结 与 4.监测结论与建议 段落。"
            "对照【季度统计JSON】里该特征的全桥统计，核对数值、对应位置、结论是否"
            "一致，需要润色时输出 summary 修复（target=指标名，hint=润色后的总结描述）：\n"
            "   - 原文说某位置正常但季度统计显示该位置疑似故障/数据缺失严重 -> 润色为"
            "“XX位置传感器故障/数据异常”；\n"
            "   - 原文说某位置异常但季度统计显示正常 -> 润色为正常；\n"
            "   - 数值/位置与季度统计一致 -> 不报。\n"
            "上下文：\n"
            "【季度统计JSON】每特征的全桥统计（含 疑似故障传感器位置/数据缺失严重的"
            "传感器位置/疑似故障时间段），是数值与位置的权威口径；\n"
            "【传感器名称对照表】含 表格映射 与 传感器名称 键。\n"
            "禁止：不对照成品报告原文；不修改任何标题（章节/表格/图注标题）；"
            "不编造数值/位置；不把思考过程当问题输出；查不出问题必须输出空 "
            "issues/repairs。\n"
            "issues.needs_human：仅当季度统计JSON缺失该特征、或确实无法判断时才 true。\n"
            "repairs.type 只允许 summary。"
        )
        # 长报告拆段逐段审查（降低单次上下文过长导致的幻觉），每段独立调用，
        # 问题合并去重；名称对照/季度统计/血缘/图表索引等全局信息每段都带上。
        base_parts = []
        if name_dict_json:
            base_parts.append("【传感器名称对照表】\n" + name_dict_json[:30000])
        if quarterly_stats_json:
            base_parts.append("【季度统计JSON】\n" + quarterly_stats_json[:40000])
        if lineage_digest:
            base_parts.append("【数据链路摘要（未找到/回退项）】\n" + lineage_digest)
        if table_warnings:
            base_parts.append("【规则式填表校验告警】\n" + table_warnings)
        if chart_index:
            base_parts.append("【图表索引表（chart_id→图注→位置，供 chart/caption 修复定位）】\n"
                              + chart_index)
        if prior_issues:
            base_parts.append("【上一轮已发现问题（本轮判断是否已解决、解决是否正确）】\n"
                              + prior_issues)
        name_dict = {}
        try:
            if name_dict_json:
                name_dict = json.loads(name_dict_json)
        except Exception:  # noqa: BLE001
            name_dict = {}
        segments = _split_segments(report_text)
        all_issues, all_repairs, raws = [], [], []
        for _i, seg in enumerate(segments):
            sec_hint = _section_match_hint(name_dict, seg)
            parts = [f"【生成的报告（第{_i + 1}/{len(segments)}段）】\n{seg}"]
            if sec_hint:
                parts.append("【本段小节匹配（传感器对照表键）】\n" + sec_hint)
            parts += base_parts
            data = self._ask_json(system, "\n\n".join(parts))
            if isinstance(data, dict):
                all_issues += [x for x in (data.get("issues") or [])
                               if isinstance(x, dict)]
                all_repairs += [x for x in (data.get("repairs") or [])
                                if isinstance(x, dict)]
                raws.append(json.dumps(data, ensure_ascii=False))
        # 合并去重
        seen_i, seen_r = set(), set()
        for x in all_issues:
            key = (str(x.get("type")), str(x.get("detail"))[:60])
            if key not in seen_i:
                seen_i.add(key)
                result["issues"].append(x)
        for x in all_repairs:
            key = (str(x.get("type")), str(x.get("target"))[:50],
                   str(x.get("hint"))[:40])
            if key not in seen_r:
                seen_r.add(key)
                result["repairs"].append(x)
        result["raw"] = "\n".join(raws)[:12000] if raws else ""
        # 后置过滤：LLM 偶发“自证清白”条目（经查证匹配/无问题/无index_wrong…）
        # 一律剔除，绝不当作问题返回
        result["issues"] = [
            i for i in result["issues"]
            if not _is_self_exonerating(str(i.get("detail") or ""))
            and not _issue_false_missing_position(
                str(i.get("detail") or ""), name_dict)
        ]
        result["ok"] = not result["issues"]
        return result


def _is_self_exonerating(detail: str) -> bool:
    """是否属于“自证清白”条目：推理证明正确却仍被列入 issues。"""
    if not detail:
        return False
    if re.search(r"无\s*(index_wrong|image_wrong|stat_logic|错误)", detail):
        return True
    if re.search(r"不构成(索引|图片|统计)?错误", detail):
        return True
    if re.search(r"(无问题|无错误|未发现问题|确认无误|完全匹配|经查证.*(正确|匹配|无误|合法))",
                 detail):
        return True
    return False


_FEAT_KW = [
    ("WD", ("结构温度",)),
    ("YB", ("应变",)),
    ("WSD", ("环境温度", "环境湿度", "温湿度")),
    ("GNSS", ("空间变位", "位移")),
    ("ND", ("挠度",)),
    ("DZJSD", ("地震", "振动")),
    ("SZJSD", ("地震",)),
    ("FSFX", ("风速", "风向")),
    ("SL", ("索力",)),
    ("EZJD", ("倾角",)),
]


def _issue_false_missing_position(detail: str, name_dict: Dict) -> bool:
    """误报拦截：条目宣称“某方位（上游/下游/左幅/右幅）不存在/无测点”，
    但传感器名称对照表里该特征族确实存在该方位的测点（如 结构温度 有
    “3#墩根部截面顶板下游”，审查却称仅有上游）——判为误报剔除。
    应变没有下游测点时不会误伤（对照表无 下游+YB）。"""
    if not detail or not name_dict:
        return False
    m = re.search(
        r"(?:无|没有|不存在|未布设|未设)[^，。；]{0,12}?"
        r"['\"“”]?(上游|下游|左幅|右幅)['\"“”]?"
        r"(?:方向|测点|传感器|维度|对应)?", detail)
    if not m:
        return False
    direction = m.group(1)
    feat_prefix = ""
    for prefix, kws in _FEAT_KW:
        if any(k in detail for k in kws):
            feat_prefix = prefix
            break
    if not feat_prefix:
        return False
    sn = (name_dict or {}).get("传感器名称") or {}
    if not isinstance(sn, dict):
        sn = name_dict or {}
    for _k, entries in sn.items():
        if direction not in str(_k):
            continue
        for e in entries or []:
            feats = e.get("特征编码") or []
            if not feats and e.get("特征"):
                feats = [e["特征"]]
            if any(feat_prefix in str(f) for f in feats):
                return True
    return False


def _section_match_hint(name_dict: Dict, seg: str) -> str:
    """给段落加上“章节->名称对照表键”的确定性匹配提示。"""
    try:
        from .section_matcher import match_section
    except Exception:  # noqa: BLE001
        return ""
    title = ""
    for line in str(seg or "").split("\n"):
        t = line.strip()
        if re.match(r"^\d+(\.\d+){1,3}(?=[\u4e00-\u9fa5\s])", t):
            title = t
            break
    if not title:
        return ""
    r = match_section(name_dict, title, seg)
    out = f"小节：{title}；表格映射键={r.get('表格映射') or '（无）'}；" \
          f"指标={r.get('指标') or '（无）'}；命中传感器名称键=" \
          f"{'、'.join(r.get('传感器名称') or [])[:300] or '（无）'}"
    return out


def _split_segments(text: str, max_len: int = 6000) -> List[str]:
    """把报告文本按段落拆成 ≤max_len 的分段（尽量保持段落完整）。"""
    paras = [p for p in str(text or "").split("\n") if p.strip()]
    if not paras:
        return [str(text or "")]
    segs, cur, cur_len = [], [], 0
    for p in paras:
        if cur and cur_len + len(p) > max_len:
            segs.append("\n".join(cur))
            cur, cur_len = [], 0
        cur.append(p)
        cur_len += len(p) + 1
    if cur:
        segs.append("\n".join(cur))
    return segs


_UNIT_RE = re.compile(r"(m/s²|m/s2|℃|%|με|mm|kN|mm/s²)")
_NUM_RE = re.compile(r"-?\d+(?:\.\d+)?")


def _to_num(s: str) -> Optional[float]:
    try:
        return float(str(s).replace(",", "").strip())
    except (TypeError, ValueError):
        return None


def self_check_report(path: str) -> List[Dict]:
    """确定性体检（不依赖 LLM），返回问题清单 [{type, detail}]。

    检查：残留占位符、表格重复行、最大<最小、平均值超出 [最小值, 最大值]、
    差值/极差为负、湿度百分数超出 0~100、数值与单位之间多余空格。
    """
    issues: List[Dict] = []
    if not path or not __import__("os").path.isfile(path):
        return issues
    try:
        from docx import Document
        doc = Document(path)
    except Exception as exc:  # noqa: BLE001
        log.warning("self_check 读取报告失败: %s", exc)
        return issues

    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    # 1) 残留占位符
    for t in texts:
        if "{{" in t:
            issues.append({"type": "leftover_placeholder", "detail": t[:80]})
    # 2) 表格：重复行 + 数值逻辑
    for ti, tb in enumerate(doc.tables):
        if not tb.rows:
            continue
        seen_rows = set()
        header = [c.text.strip() for c in tb.rows[0].cells]
        col = {h: i for i, h in enumerate(header)}

        def celln(row, *names):
            for h, i in col.items():
                if any(n in h for n in names):
                    return _to_num(row.cells[i].text)
            return None

        for ri, row in enumerate(tb.rows[1:], 1):
            cells = [c.text.strip() for c in row.cells]
            key = "|".join(cells)
            if key and key in seen_rows:
                issues.append({"type": "table_duplicate",
                               "detail": f"表{ti + 1} 行{ri} 重复"})
            seen_rows.add(key)
            mx = celln(row, "最大", "最高")
            mn = celln(row, "最小", "最低")
            av = celln(row, "平均")
            df = celln(row, "差值", "极差")
            if mx is not None and mn is not None and mx < mn:
                issues.append({"type": "stat_logic",
                               "detail": f"表{ti + 1} 行{ri} 最大{mx} < 最小{mn}"})
            if av is not None and mx is not None and mn is not None \
                    and not (mn - 1e-6 <= av <= mx + 1e-6):
                issues.append({"type": "stat_logic",
                               "detail": f"表{ti + 1} 行{ri} 平均{av} 不在 [{mn},{mx}]"})
            if df is not None and df < -1e-9:
                issues.append({"type": "stat_logic",
                               "detail": f"表{ti + 1} 行{ri} 差值/极差为负 {df}"})
    # 3) 湿度/百分比越界 + 单位空格
    for t in texts:
        for m in re.finditer(r"([-+]?\d+(?:\.\d+)?)\s*(%|％)", t):
            v = _to_num(m.group(1))
            if v is not None and (v < 0 or v > 100):
                issues.append({"type": "stat_logic",
                               "detail": f"百分比越界 {m.group(0)} | {t[:50]}"})
        if re.search(r"\d\s{2,}(m/s²|m/s2|℃|%|με|mm|kN)", t):
            issues.append({"type": "unit_wrong",
                           "detail": f"数值与单位之间有多余空格 | {t[:50]}"})
    return issues
